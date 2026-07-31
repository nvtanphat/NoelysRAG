from __future__ import annotations

import logging
import os
import re
from pathlib import Path
from typing import Any
from uuid import NAMESPACE_URL, uuid5

from src.processing import table_serializer
from src.processing.types import BBox, BlockType, DependencyUnavailableError, ParsedBlock, ParsedDocument, ParsedPage

SUPPORTED_DOCLING_EXTENSIONS = {"pdf", "pptx", "docx"}
logger = logging.getLogger(__name__)
_BULLET_PREFIX = re.compile(r"^[\u25a1\u2022\u25cf\u25aa\u25ab\u2013\u2014\-\*\u25c6\u25c7\u25cb]\s")
_SECTION_SPLIT = re.compile(r"(?=(?:^|\n)(?:\d+(?:\.\d+)*\.?\s+|[A-Z][A-Za-z ]{2,}:))")


def _workspace_cache_dir(name: str) -> Path:
    return Path(__file__).resolve().parents[3] / "data" / "cache" / name


def _is_empty_figure_only(blocks: list[ParsedBlock]) -> bool:
    """True when the page's only blocks are figure placeholders with no text.

    Docling marks scanned-text legal PDF pages as a single FIGURE block with
    `needs_captioning=True` and empty content. Those pages must still go
    through EasyOCR — otherwise the entire page text is lost downstream.
    """
    if not blocks:
        return True
    return all(
        block.block_type == BlockType.FIGURE.value and not (block.content or "").strip()
        for block in blocks
    )


def _is_garbled_vietnamese(blocks: list[ParsedBlock]) -> bool:
    """True when Docling extracted text but Vietnamese diacritics are almost absent.

    PDFs with non-standard font encoding (common in Vietnamese legal documents
    produced by older software) yield text like "QUOc HQI CQNG HOA" instead of
    "QUỐC HỘI CỘNG HÒA". The text layer exists but is garbled — OCR must run.

    Heuristic: sample the first 400 characters of page text; if < 3% are
    Vietnamese tone/diacritic characters the page is considered garbled.
    The 3% threshold is loaded from config (ocr_garbled_vi_diacritic_ratio).
    """
    import unicodedata

    text = " ".join((b.content or "") for b in blocks)[:400]
    if not text.strip():
        return False

    vi_chars = sum(
        1 for ch in text
        if unicodedata.category(ch) == "Ll" and unicodedata.combining(ch) == 0
        and ord(ch) > 127  # extended Latin — catches ă, â, ê, ô, ơ, ư + tones
    )
    total_alpha = sum(1 for ch in text if ch.isalpha())
    if total_alpha == 0:
        return False

    from src.core.config import get_settings as _get_settings
    threshold = _get_settings().ocr_garbled_vi_diacritic_ratio
    return (vi_chars / total_alpha) < threshold


def _has_vietnamese_content(blocks: list[ParsedBlock]) -> bool:
    """True when blocks contain meaningful Vietnamese diacritic text (ratio > 5%).

    Used to select VietOCR for OCR fallback pages in documents where Docling
    extracted Vietnamese text correctly from most pages but crashed on a few
    (std::bad_alloc in layout model). Those crashed pages get empty blocks →
    OCR fallback → must still use VietOCR to preserve tone marks.
    """
    import unicodedata

    text = " ".join((b.content or "") for b in blocks)[:800]
    if not text.strip():
        return False
    vi_chars = sum(
        1 for ch in text
        if unicodedata.category(ch) == "Ll" and unicodedata.combining(ch) == 0
        and ord(ch) > 127
    )
    total_alpha = sum(1 for ch in text if ch.isalpha())
    return total_alpha > 0 and (vi_chars / total_alpha) > 0.05


def _configure_docling_cache() -> None:
    os.environ.setdefault("USE_TF", "0")
    os.environ.setdefault("TRANSFORMERS_NO_TF", "1")
    os.environ.setdefault("USE_FLAX", "0")
    os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")
    cache_dirs = {
        "HF_HOME": _workspace_cache_dir("huggingface"),
        "MODELSCOPE_CACHE": _workspace_cache_dir("modelscope"),
    }
    for name, path in cache_dirs.items():
        path.mkdir(parents=True, exist_ok=True)
        os.environ.setdefault(name, str(path))
    try:
        import huggingface_hub.file_download as hf_file_download
    except ImportError:
        return
    hf_file_download.are_symlinks_supported = lambda cache_dir=None: False


def _patch_docling_transformers_compat() -> None:
    try:
        import transformers
        import transformers.image_utils as image_utils
    except ImportError:
        return
    if not hasattr(image_utils, "VideoInput"):
        setattr(image_utils, "VideoInput", getattr(image_utils, "ImageInput", object))
    if not hasattr(image_utils, "VideoMetadata"):
        setattr(image_utils, "VideoMetadata", dict)
    if hasattr(transformers, "AutoModelForImageTextToText"):
        return
    fallback = getattr(transformers, "AutoModelForVision2Seq", None)
    if fallback is not None:
        setattr(transformers, "AutoModelForImageTextToText", fallback)


class DoclingParser:
    def parse(self, file_path: Path, *, language: str = "unknown") -> ParsedDocument:
        extension = file_path.suffix.lower().lstrip(".")
        if extension not in SUPPORTED_DOCLING_EXTENSIONS:
            raise ValueError(f"DoclingParser does not support .{extension}")

        try:
            self._ensure_docling_available()
        except ImportError as exc:
            if extension == "pdf":
                logger.warning("Docling is unavailable for PDF; using pypdf/OCR fallback", extra={"error": str(exc)})
                return self._pdf_text_fallback_document(file_path, language=language, parser_error=exc)
            raise DependencyUnavailableError("docling is required for PDF/PPTX/DOCX parsing") from exc

        converter = self._converter(extension)
        try:
            result = converter.convert(str(file_path))
        except Exception as exc:
            if extension == "pdf":
                logger.warning("Docling PDF parse failed, using pypdf text fallback", extra={"error": str(exc)})
                return self._pdf_text_fallback_document(file_path, language=language, parser_error=exc)
            raise

        logger.info("Docling convert returned for %s; exporting parsed blocks", file_path.name)
        pages = self._pages_from_export(result.document, file_path=file_path, extension=extension, language=language)
        block_count = sum(len(page.blocks) for page in pages)
        logger.info("Docling export completed for %s pages=%s blocks=%s", file_path.name, len(pages), block_count)
        extra = {"parser": "docling"}
        if extension == "pdf":
            extra["pdf_strategy"] = "docling_layout_first_text_ocr_missing_pages"
        return ParsedDocument(
            source_path=str(file_path),
            file_type=extension,
            language=language,
            pages=pages,
            extra=extra,
        )

    @staticmethod
    def _ensure_docling_available() -> None:
        _configure_docling_cache()
        _patch_docling_transformers_compat()
        from docling.document_converter import DocumentConverter  # noqa: F401

    @staticmethod
    def _converter(extension: str):
        from docling.document_converter import DocumentConverter

        if extension != "pdf":
            return DocumentConverter()

        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import PdfFormatOption

        from src.core.config import get_settings as _get_settings
        _cfg = _get_settings()
        pdf_options = PdfPipelineOptions(
            do_ocr=False,
            ocr_batch_size=1,
            layout_batch_size=1,
            table_batch_size=1,
            queue_max_size=_cfg.docling_queue_max_size,
            images_scale=_cfg.docling_images_scale,
            # NOTE: formula LaTeX is taken from Docling's default text export
            # (see equation tagging in _blocks_from_docling_dict). We deliberately
            # do NOT enable do_formula_enrichment — its CodeFormulaV2 VLM crashes
            # the pipeline on CPU (→ pypdf fallback) and runs minutes/document.
            # NOTE: do NOT set generate_picture_images=True here — retaining all
            # picture pixels across the document triggers std::bad_alloc on
            # image-heavy PDFs. Figure images are instead extracted lazily,
            # page-by-page, with PyMuPDF in the captioning step (memory-light).
        )
        return DocumentConverter(format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_options)})

    def _pages_from_export(self, document: Any, *, file_path: Path, extension: str, language: str) -> list[ParsedPage]:
        logger.info("Docling export start file=%s extension=%s", file_path.name, extension)
        exported = self._export_dict(document)
        blocks = self._blocks_from_docling_dict(exported, language=language)
        logger.info("Docling export dict parsed file=%s blocks=%s", file_path.name, len(blocks))
        if extension == "docx":
            blocks.extend(self._docx_table_blocks(file_path=file_path, language=language, reading_order_offset=len(blocks)))
        if extension == "pdf":
            # Docling exports PDF tables as structured TableItem objects (no flat
            # text), so _collect_text_nodes misses them. Read document.tables and
            # re-run enrichment so the new markdown tables get HTML grid + rows.
            pdf_tables = self._pdf_table_blocks(document, language=language)
            if pdf_tables:
                blocks = self._enrich_table_blocks(blocks + pdf_tables, language=language)
            logger.info("Docling PDF table extraction file=%s tables=%s blocks=%s", file_path.name, len(pdf_tables), len(blocks))
        if not blocks:
            markdown = self._export_markdown(document)
            if markdown.strip():
                blocks = [
                    ParsedBlock(
                        block_id=self._stable_block_id(file_path, 1, 0, markdown[:80]),
                        block_index=0,
                        block_type=BlockType.PARAGRAPH.value,
                        content=markdown.strip(),
                        page_number=1,
                        language=language,
                        reading_order=0,
                        source="docling_markdown",
                    )
                ]
        pages_by_number: dict[int, list[ParsedBlock]] = {}
        for block in blocks:
            pages_by_number.setdefault(block.page_number, []).append(block)
        if extension == "pdf":
            self._add_pdf_text_fallback_pages(pages_by_number, file_path=file_path, language=language)
            self._add_easyocr_pages(pages_by_number, file_path=file_path, language=language)
        if not pages_by_number:
            return [ParsedPage(page_number=1, blocks=[])]
        page_widths = getattr(self, "_page_widths", {})
        page_heights = getattr(self, "_page_heights", {})
        return [
            ParsedPage(
                page_number=page_number,
                blocks=sorted(page_blocks, key=lambda block: block.reading_order),
                width=int(page_widths[page_number]) if page_number in page_widths else None,
                height=int(page_heights[page_number]) if page_number in page_heights else None,
            )
            for page_number, page_blocks in sorted(pages_by_number.items())
        ]

    def _pdf_text_fallback_document(self, file_path: Path, *, language: str, parser_error: Exception) -> ParsedDocument:
        pages_by_number: dict[int, list[ParsedBlock]] = {}
        self._add_pdf_text_fallback_pages(pages_by_number, file_path=file_path, language=language)
        self._add_easyocr_pages(pages_by_number, file_path=file_path, language=language)
        pages = [
            ParsedPage(page_number=page_number, blocks=sorted(page_blocks, key=lambda block: block.reading_order))
            for page_number, page_blocks in sorted(pages_by_number.items())
        ] or [ParsedPage(page_number=1, blocks=[])]
        return ParsedDocument(
            source_path=str(file_path),
            file_type="pdf",
            language=language,
            pages=pages,
            extra={"parser": "pypdf_fallback", "docling_error": f"{type(parser_error).__name__}: {parser_error}"},
        )

    def _add_pdf_text_fallback_pages(
        self,
        pages_by_number: dict[int, list[ParsedBlock]],
        *,
        file_path: Path,
        language: str,
    ) -> None:
        try:
            from pypdf import PdfReader
        except ImportError:
            return

        try:
            reader = PdfReader(str(file_path))
        except Exception:
            return

        for page_index, page in enumerate(reader.pages, start=1):
            if page_index in pages_by_number:
                continue
            try:
                text = (page.extract_text() or "").strip()
            except Exception:
                continue
            if not text:
                continue
            pages_by_number[page_index] = self._blocks_from_plain_text_page(
                text,
                file_path=file_path,
                page_number=page_index,
                language=language,
                source="pypdf_text_fallback",
                fallback_reason="docling_missing_page",
            )

    def _blocks_from_plain_text_page(
        self,
        text: str,
        *,
        file_path: Path,
        page_number: int,
        language: str,
        source: str,
        fallback_reason: str,
    ) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        for unit in self._split_plain_text_page(text):
            cleaned = unit.strip()
            if not cleaned:
                continue
            blocks.append(
                ParsedBlock(
                    block_id=self._stable_block_id(file_path, page_number, len(blocks), cleaned[:80]),
                    block_index=len(blocks),
                    block_type=self._classify_plain_text_block(cleaned),
                    content=cleaned,
                    page_number=page_number,
                    language=language,
                    reading_order=len(blocks),
                    source=source,
                    extra={"fallback_reason": fallback_reason},
                )
            )
        return blocks

    @staticmethod
    def _split_plain_text_page(text: str) -> list[str]:
        lines = [" ".join(line.split()) for line in text.splitlines() if line.strip()]
        if len(lines) <= 1:
            parts = [part.strip() for part in _SECTION_SPLIT.split(text) if part.strip()]
            return parts or [text.strip()]

        units: list[str] = []
        current: list[str] = []
        for line in lines:
            starts_new = (
                _BULLET_PREFIX.match(line) is not None
                or re.match(r"^\d+(?:\.\d+)*\.?\s+", line) is not None
                or (len(line) <= 90 and not line.endswith((".", ",", ";")))
            )
            current_is_heading = len(current) == 1 and len(current[0]) <= 90 and not current[0].endswith((".", ",", ";"))
            if current and (starts_new or current_is_heading):
                units.append(" ".join(current))
                current = []
            current.append(line)
        if current:
            units.append(" ".join(current))
        return units

    @staticmethod
    def _classify_plain_text_block(text: str) -> str:
        if "|" in text and "\n" in text:
            return BlockType.TABLE.value
        if _BULLET_PREFIX.match(text) or text.startswith(("-", "*", "\u2022")):
            return BlockType.LIST.value
        if len(text) <= 120 and not text.endswith((".", ";", ":")):
            return BlockType.HEADING.value
        return BlockType.PARAGRAPH.value

    def _add_easyocr_pages(
        self,
        pages_by_number: dict[int, list[ParsedBlock]],
        *,
        file_path: Path,
        language: str,
    ) -> None:
        if not file_path.exists():
            return
        if self._pdf_has_text_layer(file_path):
            logger.info("Skipping PDF OCR fallback for digital PDF with text layer file=%s", file_path.name)
            return
        try:
            import pypdfium2 as pdfium
        except ImportError as exc:
            raise DependencyUnavailableError("pypdfium2 is required to render scanned PDF pages for OCR") from exc

        from src.core.config import get_settings
        from src.processing.ocr_engine import EasyOCREngine, VietOCRRecognizer

        settings = get_settings()
        # Pre-scan existing blocks to determine OCR language for unknown documents:
        #  - garbled: diacritics nearly absent → font encoding broken → force VietOCR
        #  - has Vietnamese content: Docling text is fine but some pages crashed
        #    (std::bad_alloc in layout model) → OCR fallback must still use VietOCR
        _sample_blocks = [b for pg in pages_by_number.values() for b in pg]
        _has_any_text = any((b.content or "").strip() for b in _sample_blocks)
        _force_vi = language not in ("vi", "en") and (
            not _has_any_text  # fully scanned PDF — no text layer at all → default vi
            or _is_garbled_vietnamese(_sample_blocks)  # garbled font encoding
            or _has_vietnamese_content(_sample_blocks)  # good text + crashed pages
        )
        ocr_language = "vi" if (language == "vi" or _force_vi) else "en"
        output_dir = _workspace_cache_dir("pdf_page_images")
        output_dir.mkdir(parents=True, exist_ok=True)
        # Best-of-breed for Vietnamese scanned pages: EasyOCR detection + VietOCR
        # recognition (tone-accurate), honouring `ocr_recognition_engine`. The
        # dedicated image pipeline already does this; without it, EasyOCR reads
        # Vietnamese tone marks poorly ("hợp đồng" → "ợông đồng") on scanned PDFs.
        recognizer = None
        if ocr_language == "vi" and settings.ocr_recognition_engine == "vietocr":
            recognizer = VietOCRRecognizer(
                device=settings.ocr_vietocr_device,
                model_name=settings.ocr_vietocr_model_name,
            )
        engine = EasyOCREngine(lang=ocr_language, gpu=settings.ocr_easyocr_gpu, recognizer=recognizer)
        render_scale = settings.pdf_render_scale

        pdf = pdfium.PdfDocument(str(file_path))
        try:
            logger.info("PDF OCR fallback scan start file=%s pages=%s language=%s", file_path.name, len(pdf), ocr_language)
            for page_index in range(len(pdf)):
                page_number = page_index + 1
                # Skip pages Docling already parsed with REAL text content.
                # Pages where the only blocks are empty figure placeholders
                # (typical for scanned legal PDFs Docling treats as images)
                # must still go through EasyOCR — otherwise the page text is lost.
                existing = pages_by_number.get(page_number, [])
                if existing and not _is_empty_figure_only(existing):
                    # For Vietnamese docs (explicit or auto-detected), also OCR when
                    # Docling text is garbled (non-standard font → diacritics missing).
                    if ocr_language != "vi" or not _is_garbled_vietnamese(existing):
                        continue
                    # Garbled page: clear Docling blocks and re-OCR.
                    pages_by_number[page_number] = []
                logger.info("PDF OCR fallback page start file=%s page=%s", file_path.name, page_number)
                page = pdf[page_index]
                try:
                    bitmap = page.render(scale=render_scale)
                    image = bitmap.to_pil()
                    image_path = output_dir / f"{uuid5(NAMESPACE_URL, f'{file_path}:{page_number}').hex}.png"
                    image.save(image_path)
                finally:
                    page.close()

                parsed = engine.parse_image(image_path, language=language if language in {"vi", "en"} else "unknown")
                blocks = [
                    block.model_copy(update={"page_number": page_number, "block_index": index, "reading_order": index})
                    for index, block in enumerate(parsed.blocks)
                ]
                if blocks:
                    # Replace empty figure placeholder with the OCR'd text blocks.
                    pages_by_number[page_number] = blocks
                logger.info("PDF OCR fallback page done file=%s page=%s blocks=%s", file_path.name, page_number, len(blocks))
        finally:
            pdf.close()

    @staticmethod
    def _pdf_has_text_layer(file_path: Path) -> bool:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            sample = " ".join((page.extract_text() or "") for page in reader.pages[:3])
            return len(sample.strip()) > 200
        except Exception:
            return False

    @staticmethod
    def _export_dict(document: Any) -> dict[str, Any]:
        for method_name in ("export_to_dict", "model_dump", "dict"):
            method = getattr(document, method_name, None)
            if callable(method):
                data = method()
                if isinstance(data, dict):
                    return data
        return {}

    @staticmethod
    def _export_markdown(document: Any) -> str:
        method = getattr(document, "export_to_markdown", None)
        if callable(method):
            return str(method())
        return ""

    def _blocks_from_docling_dict(self, data: dict[str, Any], *, language: str) -> list[ParsedBlock]:
        candidates = self._collect_text_nodes(data)
        figure_nodes = self._collect_figure_nodes(data)
        blocks: list[ParsedBlock] = []

        # Extract page dimensions from Docling export for y-axis normalisation.
        # Docling uses PDF coordinate space: origin bottom-left, units = points.
        # We flip the y-axis so all downstream consumers use top-left origin.
        page_heights: dict[int, float] = {}
        self._page_widths: dict[int, float] = {}
        self._page_heights: dict[int, float] = {}
        for k, v in (data.get("pages") or {}).items():
            try:
                pnum = int(k)
                size = (v.get("size") or {}) if isinstance(v, dict) else {}
                w = size.get("width")
                h = size.get("height")
                if w is not None:
                    self._page_widths[pnum] = float(w)
                if h is not None:
                    page_heights[pnum] = float(h)
                    self._page_heights[pnum] = float(h)
            except (ValueError, TypeError, AttributeError):
                pass

        for index, node in enumerate(candidates):
            text = self._node_text(node)
            if not text:
                continue
            label = str(node.get("label") or node.get("type") or "").lower()
            if label in {"page_footer", "page_header"}:
                continue
            page_number = self._node_page(node)
            raw_bbox = self._node_bbox(node)
            block_type = self._classify_node(node, text)
            extra: dict[str, Any] = {"label": str(node.get("label") or node.get("type") or "")}
            # Tier 0: preserve docling's native heading depth so the mindmap can
            # build an exact hierarchy without re-deriving levels from text. A
            # plain "title" (doc title) is the shallowest; section_header carries
            # an explicit `level` (1 = top). Kept where docling provides it.
            if block_type == BlockType.HEADING.value:
                native_level = node.get("level")
                if isinstance(native_level, int):
                    extra["heading_level"] = native_level
                elif "title" in extra["label"].lower():
                    extra["heading_level"] = 0
            # Preserve the LaTeX of math formulas so the snippet stays renderable
            # and the original expression survives chunking → citation.
            if block_type == BlockType.EQUATION.value:
                extra["latex"] = text
                extra["modality"] = "equation"
            blocks.append(
                ParsedBlock(
                    block_id=node.get("self_ref") or self._stable_block_id(Path("docling"), page_number, index, text[:80]),
                    block_index=len(blocks),
                    block_type=block_type,
                    content=text,
                    page_number=page_number,
                    language=language,
                    bbox=self._flip_bbox_y(raw_bbox, page_heights.get(page_number)),
                    reading_order=len(blocks),
                    source="docling",
                    extra=extra,
                )
            )

        # Add placeholder blocks for figures/pictures that have no caption/text.
        # For DOCX/PPTX exports, Docling often exposes embedded pictures without bbox.
        for fig_index, node in enumerate(figure_nodes):
            raw_bbox = self._node_bbox(node)
            page_number = self._node_page(node)
            bbox = self._flip_bbox_y(raw_bbox, page_heights.get(page_number))
            image_meta = node.get("image") if isinstance(node.get("image"), dict) else {}
            image_uri = image_meta.get("uri") if isinstance(image_meta, dict) else None
            block_id = node.get("self_ref") or self._stable_block_id(
                Path("docling"),
                page_number,
                10000 + fig_index,
                f"figure:{page_number}:{fig_index}:{str(node.get('self_ref') or '')}",
            )
            blocks.append(
                ParsedBlock(
                    block_id=block_id,
                    block_index=len(blocks),
                    block_type=BlockType.FIGURE.value,
                    content="",  # filled by FigureCaptioner in the pipeline
                    page_number=page_number,
                    language=language,
                    bbox=bbox,
                    reading_order=len(blocks),
                    source="docling",
                    extra={
                        "label": "figure",
                        "needs_captioning": True,
                        "embedded_image_uri": image_uri,
                        "embedded_image_mimetype": image_meta.get("mimetype") if isinstance(image_meta, dict) else None,
                        "embedded_image_size": image_meta.get("size") if isinstance(image_meta, dict) else None,
                        "docling_node_ref": node.get("self_ref"),
                    },
                )
            )

        ordered = sorted(blocks, key=lambda b: (b.page_number, b.reading_order))
        return self._enrich_table_blocks(ordered, language=language)

    def _enrich_table_blocks(self, blocks: list[ParsedBlock], *, language: str) -> list[ParsedBlock]:
        """Upgrade markdown TABLE blocks to a structured HTML grid + verbalized rows.

        docling's PDF/markdown export emits pipe tables as the TABLE block content
        and no per-row blocks. We re-parse the grid, store HTML (LLM reads the
        row/column structure), and emit one verbalized paragraph per row for
        keyword/value retrieval. Already-HTML tables and non-tables pass through.
        """
        out: list[ParsedBlock] = []
        for block in blocks:
            out.append(block)
            if block.block_type != BlockType.TABLE.value:
                continue
            content = (block.content or "").strip()
            if content.startswith("<table"):
                continue  # already structured (e.g. spreadsheet/docx path)
            grid = table_serializer.parse_markdown_table(content)
            if grid is None:
                continue
            header, body = grid
            if not body:
                continue
            table_name = (block.extra.get("sheet_name") if block.extra else None) or f"table_p{block.page_number}"
            block.content = table_serializer.to_html(header, body)
            block.extra = {
                **(block.extra or {}),
                "block_kind": "table_block",
                "sheet_name": table_name,
                "columns": header,
                **table_serializer.structured_meta(header, body),
            }
            for row_index, sentence in table_serializer.verbalize_rows(
                header, body, table_name=table_name, language=language
            ):
                out.append(
                    ParsedBlock(
                        block_id=self._stable_block_id(
                            Path("docling"), block.page_number,
                            30000 + row_index, sentence[:60],
                        ),
                        block_index=len(out),
                        block_type=BlockType.PARAGRAPH.value,
                        content=sentence,
                        page_number=block.page_number,
                        language=language,
                        bbox=block.bbox,
                        reading_order=block.reading_order,
                        source="docling",
                        extra={
                            "block_kind": "table_row",
                            "sheet_name": table_name,
                            "row_index": row_index,
                            "columns": header,
                        },
                    )
                )
        return out

    def _docx_table_blocks(self, *, file_path: Path, language: str, reading_order_offset: int) -> list[ParsedBlock]:
        try:
            from docx import Document
        except ImportError:
            logger.debug("python-docx is unavailable; skipping DOCX table augmentation")
            return []
        try:
            document = Document(str(file_path))
        except Exception as exc:
            logger.debug("Could not inspect DOCX tables", extra={"path": str(file_path), "error": str(exc)})
            return []

        blocks: list[ParsedBlock] = []
        for table_index, table in enumerate(document.tables):
            rows = [
                [self._normalize_cell_text(cell.text) for cell in row.cells]
                for row in table.rows
            ]
            rows = [row for row in rows if any(cell for cell in row)]
            if len(rows) < 2 or max((len(row) for row in rows), default=0) < 2:
                continue
            header, body = rows[0], rows[1:]
            grid_html = table_serializer.to_html(header, body)
            if not grid_html:
                continue
            table_name = f"docx_table_{table_index + 1}"
            blocks.append(
                ParsedBlock(
                    block_id=self._stable_block_id(file_path, 1, 20000 + table_index, grid_html[:80]),
                    block_index=reading_order_offset + len(blocks),
                    block_type=BlockType.TABLE.value,
                    content=grid_html,
                    page_number=1,
                    language=language,
                    reading_order=reading_order_offset + len(blocks),
                    source="python_docx",
                    extra={
                        "label": "docx_table",
                        "table_index": table_index,
                        "block_kind": "table_block",
                        "sheet_name": table_name,
                        "columns": header,
                        **table_serializer.structured_meta(header, body),
                    },
                )
            )
            # Verbalized rows → keyword/value retrieval (PDF/DOCX tables had none)
            for row_index, sentence in table_serializer.verbalize_rows(
                header, body, table_name=table_name, language=language
            ):
                blocks.append(
                    ParsedBlock(
                        block_id=self._stable_block_id(
                            file_path, 1, 21000 + table_index * 100 + row_index, sentence[:60]
                        ),
                        block_index=reading_order_offset + len(blocks),
                        block_type=BlockType.PARAGRAPH.value,
                        content=sentence,
                        page_number=1,
                        language=language,
                        reading_order=reading_order_offset + len(blocks),
                        source="python_docx",
                        extra={
                            "block_kind": "table_row",
                            "sheet_name": table_name,
                            "row_index": row_index,
                            "columns": header,
                        },
                    )
                )
        return blocks

    def _pdf_table_blocks(self, document: Any, *, language: str) -> list[ParsedBlock]:
        """Extract Docling TableFormer tables (PDF) as markdown TABLE blocks.

        Docling exports PDF tables as structured ``TableItem`` objects with a cell
        grid and no flat text, so ``_collect_text_nodes`` never picks them up. We
        read ``document.tables`` directly. Defensive: a per-table failure is
        skipped (logged), never crashes the parse. The resulting markdown blocks
        flow through ``_enrich_table_blocks`` → HTML grid + verbalized rows.
        """
        tables = getattr(document, "tables", None)
        if not tables:
            return []
        page_heights = getattr(self, "_page_heights", {})
        blocks: list[ParsedBlock] = []
        for table_index, table in enumerate(tables):
            try:
                header, rows = self._table_grid(table, document)
            except Exception as exc:
                logger.debug(
                    "PDF table extraction skipped",
                    extra={"table_index": table_index, "error": str(exc)},
                )
                continue
            if not header or not rows:
                continue
            page_number, raw_bbox = self._table_prov(table)
            content = table_serializer.to_markdown(header, rows)
            blocks.append(
                ParsedBlock(
                    block_id=getattr(table, "self_ref", None)
                    or self._stable_block_id(Path("docling"), page_number, 25000 + table_index, content[:80]),
                    block_index=len(blocks),
                    block_type=BlockType.TABLE.value,
                    content=content,
                    page_number=page_number,
                    language=language,
                    bbox=self._flip_bbox_y(raw_bbox, page_heights.get(page_number)),
                    # Place after text blocks on the page so reading order stays sane.
                    reading_order=100000 + table_index,
                    source="docling",
                    extra={"label": "table", "table_source": "docling", "table_index": table_index},
                )
            )
        return blocks

    @staticmethod
    def _clean_table_header(columns: Any) -> list[str]:
        """Collapse a 2-level (MultiIndex) PDF table header to the real leaf names.

        Docling exports tables with a spanning super-header (e.g. "Huấn luyện mô
        hình theo truyền thống") flattened onto every column, yielding garbage like
        ``"Mô hình.Huấn luyện…"``, ``"Tập kiểm tra.Huấn luyện…"``. We keep the part
        that VARIES across columns (the true header) and drop the shared span.
        """
        # MultiIndex tuples → take the most specific non-empty level per column.
        try:
            import pandas as pd  # docling dependency
            if isinstance(columns, pd.MultiIndex):
                out: list[str] = []
                for tup in columns:
                    parts = [str(p).strip() for p in tup
                             if str(p).strip() and not str(p).startswith("Unnamed")]
                    out.append(parts[-1] if parts else "")
                return out
        except Exception:
            pass

        from collections import Counter as _Counter
        cols = [str(c).strip() for c in columns]

        def _looks_like_span(text: str) -> bool:
            # A real spanning header is a phrase/unit, not a decimal fragment like "0".
            return " " in text or len(text) >= 5

        # A spanning super-header gets joined as "<leaf>.<span>" (or "<span>.<leaf>")
        # onto the columns it covers — possibly not every column (row-label cols stay
        # clean). Strip the component shared by >=2 dotted columns, keep the leaf.
        dotted = [(c.split(".", 1)[0].strip(), c.split(".", 1)[1].strip()) for c in cols if "." in c]
        if len(dotted) >= 2:
            suf_common, suf_n = _Counter(s for _, s in dotted).most_common(1)[0]
            pre_common, pre_n = _Counter(p for p, _ in dotted).most_common(1)[0]
            if suf_n >= 2 and suf_n == len(dotted) and _looks_like_span(suf_common):
                return [c.split(".", 1)[0].strip()
                        if "." in c and c.split(".", 1)[1].strip() == suf_common else c
                        for c in cols]
            if pre_n >= 2 and pre_n == len(dotted) and _looks_like_span(pre_common):
                return [c.split(".", 1)[1].strip()
                        if "." in c and c.split(".", 1)[0].strip() == pre_common else c
                        for c in cols]
        return cols

    @staticmethod
    def _table_grid(table: Any, document: Any) -> tuple[list[str], list[list[str]]]:
        """Recover ``(header, rows)`` from a Docling TableItem, version-tolerant.

        Tries, in order: pandas dataframe export, the raw ``data.grid`` cell
        matrix, then a markdown export parsed by table_serializer. Returns
        ``([], [])`` when nothing usable is found.
        """
        # 1) pandas dataframe — most stable across Docling versions.
        # Newer Docling wants the doc arg; older takes none — try doc first.
        to_df = getattr(table, "export_to_dataframe", None)
        if callable(to_df):
            try:
                df = to_df(document)
            except TypeError:
                df = to_df()
            if df is not None and not df.empty:
                header = DoclingParser._clean_table_header(df.columns)
                rows = [["" if v is None else str(v) for v in row] for row in df.values.tolist()]
                if header and rows:
                    return header, rows

        # 2) raw cell grid
        data = getattr(table, "data", None)
        grid = getattr(data, "grid", None) if data is not None else None
        if grid:
            matrix = [
                [(getattr(cell, "text", "") or "").strip() for cell in grid_row]
                for grid_row in grid
            ]
            matrix = [r for r in matrix if any(r)]
            if len(matrix) >= 2 and max((len(r) for r in matrix), default=0) >= 2:
                return matrix[0], matrix[1:]

        # 3) markdown export
        to_md = getattr(table, "export_to_markdown", None)
        if callable(to_md):
            try:
                md = to_md(document)
            except TypeError:
                md = to_md()
            parsed = table_serializer.parse_markdown_table(md or "")
            if parsed is not None:
                return parsed
        return [], []

    @staticmethod
    def _table_prov(table: Any) -> tuple[int, BBox | None]:
        """Page number + PDF-space bbox (origin bottom-left) for a TableItem."""
        prov = getattr(table, "prov", None)
        if not prov:
            return 1, None
        first = prov[0]
        page = getattr(first, "page_no", None) or 1
        bb = getattr(first, "bbox", None)
        bbox: BBox | None = None
        if bb is not None:
            left = getattr(bb, "l", None)
            top = getattr(bb, "t", None)
            right = getattr(bb, "r", None)
            bottom = getattr(bb, "b", None)
            if all(v is not None for v in (left, top, right, bottom)):
                bbox = BBox(x1=float(left), y1=float(top), x2=float(right), y2=float(bottom))
        return int(page), bbox

    @staticmethod
    def _normalize_cell_text(text: str) -> str:
        return " ".join(text.split())

    @staticmethod
    def _rows_to_markdown_table(rows: list[list[str]]) -> str:
        width = max((len(row) for row in rows), default=0)
        if width < 2:
            return ""
        padded = [(row + [""] * width)[:width] for row in rows]
        header, body = padded[0], padded[1:]
        separator = ["---"] * width

        def render(row: list[str]) -> str:
            return "| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |"

        return "\n".join([render(header), render(separator), *(render(row) for row in body)])

    def _collect_figure_nodes(self, value: Any) -> list[dict[str, Any]]:
        """Collect figure/picture nodes that have no text.

        Docling may export embedded DOCX pictures without bbox, but the node still
        carries image metadata and can be captioned downstream.
        """
        nodes: list[dict[str, Any]] = []
        if isinstance(value, dict):
            label = str(value.get("label") or value.get("type") or "").lower()
            has_text = bool(self._node_text(value))
            has_image = isinstance(value.get("image"), dict) and bool(value["image"].get("uri"))
            if ("picture" in label or "figure" in label) and not has_text and (has_image or self._node_bbox(value) is not None):
                nodes.append(value)
            for child in value.values():
                nodes.extend(self._collect_figure_nodes(child))
        elif isinstance(value, list):
            for item in value:
                nodes.extend(self._collect_figure_nodes(item))
        return nodes

    def _collect_text_nodes(self, value: Any) -> list[dict[str, Any]]:
        nodes: list[dict[str, Any]] = []
        if isinstance(value, dict):
            if self._node_text(value):
                nodes.append(value)
            for child in value.values():
                nodes.extend(self._collect_text_nodes(child))
        elif isinstance(value, list):
            for item in value:
                nodes.extend(self._collect_text_nodes(item))
        return nodes

    @staticmethod
    def _node_text(node: dict[str, Any]) -> str:
        for key in ("text", "orig", "content"):
            value = node.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
        return ""

    @staticmethod
    def _node_page(node: dict[str, Any]) -> int:
        prov = node.get("prov")
        if isinstance(prov, list) and prov:
            page_no = prov[0].get("page_no") if isinstance(prov[0], dict) else None
            if isinstance(page_no, int):
                return page_no
        page_no = node.get("page_no") or node.get("page")
        return int(page_no) if isinstance(page_no, int | float | str) and str(page_no).isdigit() else 1

    @staticmethod
    def _node_bbox(node: dict[str, Any]) -> BBox | None:
        prov = node.get("prov")
        candidate = None
        if isinstance(prov, list) and prov and isinstance(prov[0], dict):
            candidate = prov[0].get("bbox")
        candidate = candidate or node.get("bbox")
        if isinstance(candidate, dict):
            left = candidate.get("l", candidate.get("left", candidate.get("x1")))
            top = candidate.get("t", candidate.get("top", candidate.get("y1")))
            right = candidate.get("r", candidate.get("right", candidate.get("x2")))
            bottom = candidate.get("b", candidate.get("bottom", candidate.get("y2")))
            if all(value is not None for value in (left, top, right, bottom)):
                return BBox(x1=float(left), y1=float(top), x2=float(right), y2=float(bottom))
        return None

    @staticmethod
    def _flip_bbox_y(bbox: BBox | None, page_height: float | None) -> BBox | None:
        """Convert Docling PDF bbox (origin bottom-left) to screen coords (origin top-left).

        Docling's prov bbox uses `t` (top) and `b` (bottom) measured from the page bottom
        in points. We store them as y1=t, y2=b, so after the flip:
          y1_screen = page_height - bbox.y1   (top edge in screen space)
          y2_screen = page_height - bbox.y2   (bottom edge in screen space)
        Result always has y1 < y2, consistent with EasyOCR pixel coords.
        """
        if bbox is None or page_height is None or page_height <= 0:
            return bbox
        return BBox(x1=bbox.x1, y1=page_height - bbox.y1, x2=bbox.x2, y2=page_height - bbox.y2)

    @staticmethod
    def _classify_node(node: dict[str, Any], text: str) -> str:
        label = str(node.get("label") or node.get("type") or "").lower()
        if "table" in label:
            return BlockType.TABLE.value
        if "formula" in label or "equation" in label:
            return BlockType.EQUATION.value
        if "list" in label:
            return BlockType.LIST.value
        if "picture" in label or "figure" in label:
            return BlockType.FIGURE.value
        if "section_header" in label or "heading" in label or "title" in label:
            return BlockType.HEADING.value
        if text.strip().startswith(("-", "*", "\u2022")):
            return BlockType.LIST.value
        return BlockType.PARAGRAPH.value

    @staticmethod
    def _stable_block_id(file_path: Path, page_number: int, index: int, seed: str) -> str:
        return f"blk-{uuid5(NAMESPACE_URL, f'{file_path}:{page_number}:{index}:{seed}').hex[:12]}"
